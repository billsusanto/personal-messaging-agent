from pathlib import Path

from docx import Document as DocxDocument
from pydantic import BaseModel
from pypdf import PdfReader


class Document(BaseModel):
    content: str
    metadata: dict[str, str] = {}


CHUNK_SIZE = 500
CHUNK_OVERLAP = 50


def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    if not text.strip():
        return []

    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = start + chunk_size
        chunk = text[start:end]

        if chunk.strip():
            chunks.append(chunk.strip())

        start = end - overlap
        if start >= text_length:
            break

    return chunks


def load_pdf(file_path: str) -> list[str]:
    reader = PdfReader(file_path)
    full_text = ""

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            full_text += page_text + "\n"

    return _chunk_text(full_text)


def load_docx(file_path: str) -> list[str]:
    doc = DocxDocument(file_path)
    full_text = ""

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text += paragraph.text + "\n"

    return _chunk_text(full_text)


def load_documents(docs_dir: str) -> list[Document]:
    docs_path = Path(docs_dir)
    documents: list[Document] = []

    if not docs_path.exists():
        return documents

    for file_path in docs_path.iterdir():
        if not file_path.is_file():
            continue

        suffix = file_path.suffix.lower()
        chunks: list[str] = []

        if suffix == ".pdf":
            chunks = load_pdf(str(file_path))
        elif suffix in (".docx", ".doc"):
            chunks = load_docx(str(file_path))
        else:
            continue

        for i, chunk in enumerate(chunks):
            documents.append(
                Document(
                    content=chunk,
                    metadata={
                        "source": file_path.name,
                        "chunk_index": str(i),
                    },
                )
            )

    return documents
